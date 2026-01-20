import os
import io
import json
from datetime import date

import streamlit as st
from dotenv import load_dotenv
from docx import Document

from google import genai
from google.genai.types import GenerateContentConfig

import db
from prompts import JOB_EXTRACT_SYSTEM, TAILOR_SYSTEM

# -------------------- Setup --------------------
load_dotenv()
st.set_page_config(page_title="Executive Job Tracker + Resume Tailor", layout="wide")

api_key = (os.getenv("GEMINI_API_KEY") or "").strip()
if not api_key:
    st.error("Missing GEMINI_API_KEY. Add it to your .env file.")
    st.stop()

client = genai.Client(api_key=api_key)
db.init_db()

TOOLS_URL_CONTEXT = [{"url_context": {}}]

STATUSES = [
    "Target",
    "Applied",
    "Recruiter Screen",
    "Hiring Manager",
    "Panel",
    "Final Round",
    "Offer",
    "Finished",   # replaces Rejected
    "Paused",
]

FINISHED_OUTCOMES = ["", "Not selected", "Withdrew", "Role closed", "Offer declined", "Offer accepted"]

DATE_FMT = "MM/DD/YYYY"  # display format in UI


# -------------------- Session State --------------------
if "last_generation" not in st.session_state:
    st.session_state.last_generation = None
if "last_saved_message" not in st.session_state:
    st.session_state.last_saved_message = ""
if "prev_job_pick" not in st.session_state:
    st.session_state.prev_job_pick = None
if "prev_resume_pick" not in st.session_state:
    st.session_state.prev_resume_pick = None


# -------------------- Helpers --------------------
def read_docx(uploaded_file) -> str:
    doc = Document(uploaded_file)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts).strip()

def docx_bytes(title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def safe_json(text: str) -> dict:
    if text is None:
        raise ValueError("Empty model response (no text).")
    t = text.strip()
    if t.startswith("```"):
        t = t.strip("`")
        t = t.replace("json\n", "", 1).strip()
    return json.loads(t)

def iso_or_empty(d):
    return d.isoformat() if d else ""

def iso_to_date_or_none(s):
    if not s:
        return None
    try:
        return date.fromisoformat(s)
    except Exception:
        return None


# -------------------- Gemini Calls --------------------
def extract_job_from_url(url: str) -> dict:
    prompt = f"""
Read the job posting at this URL and extract the details.
URL: {url}

Return JSON only with keys: company, title, location, jd_text
"""
    resp = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[JOB_EXTRACT_SYSTEM, prompt],
        config=GenerateContentConfig(tools=TOOLS_URL_CONTEXT),
    )

    text = getattr(resp, "text", None)
    if isinstance(text, str) and text.strip():
        return safe_json(text)

    # fallback attempt (some SDK variants)
    try:
        candidates = getattr(resp, "candidates", None) or []
        if candidates:
            content = getattr(candidates[0], "content", None)
            parts = getattr(content, "parts", None) or []
            joined = "\n".join([getattr(p, "text", "") for p in parts]).strip()
            if joined:
                return safe_json(joined)
    except Exception:
        pass

    raise RuntimeError("URL could not be read (blocked/paywalled). Paste the JD manually.")

def tailor_resume(base_resume_text: str, jd_text: str) -> dict:
    prompt = f"""
Base resume (truth source):
\"\"\"\n{base_resume_text}\n\"\"\"\n
Job description:
\"\"\"\n{jd_text}\n\"\"\"\n

Return valid JSON ONLY with keys:
tailored_resume (string),
changes_summary (array of strings),
suggested_additions (array of strings),
accuracy_checklist (array of strings)
"""
    resp = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[TAILOR_SYSTEM, prompt],
    )
    return safe_json(getattr(resp, "text", "") or "")


# -------------------- UI --------------------
st.title("Executive Job Tracker + Resume Tailor")

# -------- User context (multi-user ready) --------

st.sidebar.header("User")

current_user = st.sidebar.text_input(
    "Current user email",
    value="ajay@dmsapiens.com"
).strip().lower()

if not current_user:
    st.stop()

st.session_state["current_user_email"] = current_user

# -----------------------------------------------

tabs_top = st.tabs(["Tracker", "Reports"])


# =========================================================
# TAB 1: TRACKER
# =========================================================
with tabs_top[0]:
    left, right = st.columns([1, 2])

    # ---------------- LEFT ----------------
    with left:
        st.subheader("1) Base Resume")
        uploaded = st.file_uploader("Upload base resume (.docx)", type=["docx"])
        resume_name = st.text_input("Resume name (example: Exec_Engineering_Leadership)")

        if uploaded and resume_name:
            resume_text = read_docx(uploaded)
            if st.button("Save base resume"):
                db.insert_resume(resume_name.strip(), resume_text)
                st.success("Resume saved")

        resumes = db.list_resumes()
        if resumes:
            resume_pick = st.selectbox(
                "Select base resume",
                options=[r["id"] for r in resumes],
                format_func=lambda rid: db.get_resume(rid)["name"],
            )
        else:
            resume_pick = None
            st.info("Upload a resume first")

        st.divider()
        st.subheader("2) Add Job")

        job_name = st.text_input("Job name (required) — example: Google – Director of Engineering")
        job_url = st.text_input("Job posting URL (optional)")
        manual_jd = st.text_area("If URL fails, paste job description here (optional)", height=160)

        if st.button("Save job (URL or pasted JD)"):
            try:
                if not job_name.strip():
                    st.error("Please enter a job name.")
                    st.stop()

                jd_text = ""
                job_meta = {"title": "", "location": ""}

                if job_url.strip():
                    with st.spinner("Reading job URL..."):
                        job_data = extract_job_from_url(job_url.strip())
                        jd_text = (job_data.get("jd_text") or "").strip()
                        job_meta["title"] = (job_data.get("title") or "").strip()
                        job_meta["location"] = (job_data.get("location") or "").strip()

                if not jd_text:
                    jd_text = (manual_jd or "").strip()

                if not jd_text:
                    st.error("Provide a working URL OR paste the job description text.")
                else:
                    job_id = db.insert_job(
                        company=job_name.strip(),
                        title=job_meta["title"],
                        location=job_meta["location"],
                        url=job_url.strip() if job_url.strip() else "(pasted)",
                        jd_text=jd_text,
                    )
                    st.success(f"Job saved (ID {job_id})")
            except Exception as e:
                jd_text = (manual_jd or "").strip()
                if job_name.strip() and jd_text:
                    job_id = db.insert_job(
                        company=job_name.strip(),
                        title="",
                        location="",
                        url=job_url.strip() if job_url.strip() else "(pasted)",
                        jd_text=jd_text,
                    )
                    st.success(f"Job saved using pasted JD (ID {job_id})")
                else:
                    st.error(f"Failed to save job. Details: {e}")

        st.divider()
        st.subheader("3) Select Job")

        jobs = db.list_jobs()
        if jobs:
            job_pick = st.selectbox(
                "Choose job",
                options=[j["id"] for j in jobs],
                format_func=lambda jid: f"{db.get_job(jid)['status']} | {db.get_job(jid)['company']}",
            )

            job_row = db.get_job(job_pick)

            # Only clear generated content if job/resume actually changed
            if (st.session_state.prev_job_pick != job_pick) or (st.session_state.prev_resume_pick != resume_pick):
                st.session_state.last_generation = None
                st.session_state.last_saved_message = ""
                st.session_state.prev_job_pick = job_pick
                st.session_state.prev_resume_pick = resume_pick

            new_status = st.selectbox(
                "Status",
                STATUSES,
                index=STATUSES.index(job_row["status"]) if job_row["status"] in STATUSES else 0,
            )

            status_date_val = None
            if new_status != "Target":
                status_date_val = st.date_input(
                    f"{new_status} date (required)",
                    value=iso_to_date_or_none(job_row["status_date"]) or date.today(),
                    format=DATE_FMT,
                )

            if st.button("Update status"):
                if new_status != "Target" and not status_date_val:
                    st.error("Please select a date for this status.")
                else:
                    db.update_job_status(
                        job_id=job_pick,
                        status=new_status,
                        status_date=iso_or_empty(status_date_val),
                    )
                    st.success("Status updated")

            st.markdown("### Dates + Follow-up")

            applied_date = st.date_input(
                "Applied Date (optional)",
                value=iso_to_date_or_none(job_row["applied_date"]),
                format=DATE_FMT,
            )
            followup_date = st.date_input(
                "Follow Up Date (optional)",
                value=iso_to_date_or_none(job_row["followup_date"]),
                format=DATE_FMT,
            )
            finished_date = st.date_input(
                "Finished Date (optional)",
                value=iso_to_date_or_none(job_row["finished_date"]),
                format=DATE_FMT,
            )

            finished_outcome = st.selectbox(
                "Finished Outcome (optional)",
                FINISHED_OUTCOMES,
                index=FINISHED_OUTCOMES.index(job_row["finished_outcome"]) if job_row["finished_outcome"] in FINISHED_OUTCOMES else 0,
            )

            notes = st.text_area("Notes (optional)", value=job_row["notes"] or "", height=100)

            if st.button("Save dates/notes"):
                db.update_job_dates(
                    job_id=job_pick,
                    applied_date=iso_or_empty(applied_date),
                    followup_date=iso_or_empty(followup_date),
                    finished_date=iso_or_empty(finished_date),
                    finished_outcome=finished_outcome or "",
                    notes=notes or "",
                )
                st.success("Saved dates/notes")

        else:
            job_pick = None
            st.info("Add a job first")

    # ---------------- RIGHT ----------------
    with right:
        if not (resume_pick and job_pick):
            st.info("Select a base resume + a job to generate a tailored version")
            st.stop()

        base = db.get_resume(resume_pick)
        job = db.get_job(job_pick)

        st.subheader("Job Details")
        st.write("Job name:", job["company"])
        if job["title"]:
            st.write("Title (if extracted):", job["title"])
        if job["location"]:
            st.write("Location (if extracted):", job["location"])
        st.write("URL:", job["url"])

        with st.expander("View Job Description"):
            st.text_area("JD", value=job["jd_text"], height=220)

        st.divider()
        st.subheader("Generate Tailored Resume")

        default_name = f"{job['company'].replace(' ', '_')}_v1"
        version_name = st.text_input("Version name", value=default_name)

        if st.button("Generate", type="primary"):
            st.session_state.last_saved_message = ""
            with st.spinner("Generating tailored resume..."):
                result = tailor_resume(base["resume_text"], job["jd_text"])

            tailored = result.get("tailored_resume", "") or ""
            changes = result.get("changes_summary", []) or []
            additions = result.get("suggested_additions", []) or []
            checklist = result.get("accuracy_checklist", []) or []

            if not tailored:
                st.error("Model did not include 'tailored_resume'. Raw response below:")
                st.json(result)
                st.session_state.last_generation = None
            else:
                st.session_state.last_generation = {
                    "version_name": (version_name or "").strip() or "version",
                    "tailored_resume": tailored,
                    "changes_summary": changes,
                    "suggested_additions": additions,
                    "accuracy_checklist": checklist,
                }

        gen = st.session_state.last_generation
        if gen:
            tabs = st.tabs(["Resume", "Changes", "Suggestions", "Checklist"])

            with tabs[0]:
                st.text_area("Tailored Resume", gen["tailored_resume"], height=400)
                st.download_button(
                    "Download DOCX",
                    data=docx_bytes("Tailored Resume", gen["tailored_resume"]),
                    file_name=f"{gen['version_name']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            with tabs[1]:
                for x in gen["changes_summary"] or ["(none returned)"]:
                    st.write("•", x)

            with tabs[2]:
                for x in gen["suggested_additions"] or ["(none returned)"]:
                    st.write("•", x)

            with tabs[3]:
                for x in gen["accuracy_checklist"] or ["(none returned)"]:
                    st.write("☐", x)

            # ✅ Save should NOT clear the screen; gen stays in session_state
            if st.button("Save Version to Job"):
                db.insert_version(
                    job_id=job_pick,
                    base_resume_id=resume_pick,
                    version_name=gen["version_name"],
                    tailored_resume=gen["tailored_resume"],
                    changes_summary=json.dumps(gen["changes_summary"]),
                    suggested_additions=json.dumps(gen["suggested_additions"]),
                    accuracy_checklist=json.dumps(gen["accuracy_checklist"]),
                )
                st.session_state.last_saved_message = f"Saved: {gen['version_name']}"
                st.success(st.session_state.last_saved_message)

        st.divider()
        st.subheader("Saved Versions (shows base resume used)")

        versions = db.list_versions_for_job(job_pick)
        if versions:
            for v in versions:
                header = f"{v['version_name']}  |  Base resume: {v['resume_name']}  |  {v['created_at']}"
                with st.expander(header):
                    st.text_area("Resume", v["tailored_resume"], height=220)
        else:
            st.write("No versions saved yet")


# =========================================================
# TAB 2: REPORTS
# =========================================================
with tabs_top[1]:
    st.subheader("Reports")

    today = date.today().isoformat()

    st.markdown("### Follow-ups due (today or earlier)")
    due = db.followups_due_rows(today)
    if due:
        st.dataframe(due, use_container_width=True)
    else:
        st.write("No follow-ups due.")

    st.divider()
    st.markdown("### Pipeline by status")
    rows = db.jobs_report_rows()

    counts = {}
    for r in rows:
        s = r["status"] or "Target"
        counts[s] = counts.get(s, 0) + 1

    import pandas as pd
    counts_df = pd.DataFrame([{"Status": k, "Count": v} for k, v in counts.items()]).sort_values(
        ["Count", "Status"], ascending=[False, True]
    )
    st.dataframe(counts_df, use_container_width=True)

    st.divider()
    st.markdown("### Full jobs table (export)")

    df = pd.DataFrame([dict(r) for r in rows])

    # Friendly column names for display/export
    rename_map = {
        "id": "ID",
        "company": "Job Name",
        "status": "Status",
        "status_date": "Status Date",
        "applied_date": "Applied Date",
        "followup_date": "Follow Up Date",
        "finished_date": "Finished Date",
        "finished_outcome": "Finished Outcome",
        "notes": "Notes",
        "created_at": "Created At",
    }
    df = df.rename(columns=rename_map)

    # On-screen with bold headers (where supported)
    try:
        styled = df.style.set_table_styles([{"selector": "th", "props": [("font-weight", "bold")]}])
        st.dataframe(styled, use_container_width=True)
    except Exception:
        st.dataframe(df, use_container_width=True)

    # CSV download (no formatting, but friendly headers)
    st.download_button(
        "Download jobs report as CSV",
        data=df.to_csv(index=False).encode("utf-8"),
        file_name="jobs_report.csv",
        mime="text/csv",
    )

    # Excel download with bold header row
    try:
        import openpyxl
        from openpyxl.styles import Font

        out = io.BytesIO()
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Jobs Report"

        headers = list(df.columns)
        ws.append(headers)
        for c in range(1, len(headers) + 1):
            ws.cell(row=1, column=c).font = Font(bold=True)

        for row in df.itertuples(index=False):
            ws.append(list(row))

        wb.save(out)
        out.seek(0)

        st.download_button(
            "Download jobs report as Excel (bold headers)",
            data=out.getvalue(),
            file_name="jobs_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception:
        st.info("Excel export needs openpyxl. Install with:  py -3.12 -m pip install openpyxl")

